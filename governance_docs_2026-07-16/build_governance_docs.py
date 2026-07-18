from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
GRAY = "555555"
LIGHT_GRAY = "F2F4F7"
BORDER = "D7DCE2"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color="000000", bold=None, italic=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), name)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_page_field(paragraph):
    run = paragraph.add_run("Стр. ")
    set_run_font(run, size=9, color=GRAY)
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=GRAY)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), "Calibri")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def _format_base_run(run, base_size=None, base_color=None, bold=False):
    if base_size is None and base_color is None and not bold:
        return
    if base_size is None and base_color is None:
        run.bold = bold
        return
    set_run_font(run, size=base_size, color=base_color, bold=bold)


def add_inline(paragraph, text, *, base_size=None, base_color=None, bold=False):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            _format_base_run(run, base_size, base_color, bold)
        token = match.group(0)
        if token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _format_base_run(run, base_size, base_color, True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            code_size = max(9, (base_size or 11) - 0.5)
            set_run_font(run, name="Consolas", size=code_size, color=DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _format_base_run(run, base_size, base_color, bold)


def paragraph_border_bottom(paragraph, color=BLUE, size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def create_num_id(doc, kind="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)

    for level in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        if kind == "decimal":
            lvl_text.set(qn("w:val"), f"%{level + 1}." if level == 0 else f"%{level + 1})")
        else:
            lvl_text.set(qn("w:val"), "•" if level % 2 == 0 else "–")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(720 + level * 540))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 540))
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def choose_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 2:
        second_numeric = all(re.fullmatch(r"[\d\s–—/.,%]+", r[1].strip()) for r in rows[1:] if len(r) > 1)
        return [7800, 1560] if second_numeric else [2700, 6660]
    if cols == 3:
        return [2160, 3600, 3600]
    if cols == 4:
        return [1450, 2200, 2500, 3210]
    base = CONTENT_DXA // cols
    widths = [base] * cols
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def parse_md_table(lines: list[str], start: int):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if c_idx > 0 and len(text) < 16:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, text.replace("**", ""), base_size=9.5, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, LIGHT_GRAY)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Calibri", 25, NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 12.5, GRAY, italic=False)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Document Metadata" not in styles:
        meta = styles.add_style("Document Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Document Metadata"]
    set_style_font(meta, "Calibri", 10.5, GRAY)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.line_spacing = 1.10

    if "Formula Block" not in styles:
        formula = styles.add_style("Formula Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = styles["Formula Block"]
    set_style_font(formula, "Consolas", 10, DARK_BLUE)
    formula.paragraph_format.left_indent = Inches(0.25)
    formula.paragraph_format.right_indent = Inches(0.15)
    formula.paragraph_format.space_before = Pt(3)
    formula.paragraph_format.space_after = Pt(8)
    formula.paragraph_format.keep_together = True


def configure_section(doc, running_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run(running_label)
    set_run_font(left, size=8.5, color=GRAY, bold=True)
    right = p.add_run("\tПроект 0.2 • 16.07.2026")
    set_run_font(right, size=8.5, color=GRAY)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    for footer in (section.footer, section.first_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        add_page_field(p)


def split_front_matter(lines):
    title = lines[0].lstrip("# ").strip()
    meta = []
    idx = 1
    while idx < len(lines) and not lines[idx].startswith("## ") and not lines[idx].startswith("# "):
        if lines[idx].strip():
            meta.append(lines[idx].strip())
        idx += 1
    return title, meta, idx


def add_memo_masthead(doc, title, meta):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("ПРОЕКТ НОРМАТИВНОГО ДОКУМЕНТА")
    set_run_font(run, size=9, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.keep_with_next = True
    add_inline(p, title, base_size=25, base_color=NAVY, bold=True)
    for line in meta:
        p = doc.add_paragraph(style="Document Metadata")
        add_inline(p, line, base_size=10.5, base_color=GRAY)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(10)
    paragraph_border_bottom(rule, BLUE, "8", "4")


def add_editorial_cover(doc, title, meta):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("ИНСТИТУЦИОНАЛЬНОЕ ОБОСНОВАНИЕ")
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_inline(p, title, base_size=26, base_color=NAVY, bold=True)

    for line in meta:
        p = doc.add_paragraph(style="Document Metadata")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line, base_size=10.5, base_color=GRAY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Экспертный вес без потолка • самостоятельная образовательная надбавка • усиленный Госплан")
    set_run_font(run, size=10, color=DARK_BLUE, italic=True)
    p.add_run().add_break(WD_BREAK.PAGE)


LIST_RE = re.compile(r"^(\s*)([-*]|\d+\.)\s+(.*)$")


def add_markdown_body(doc, lines, start_idx):
    idx = start_idx
    current_list_kind = None
    current_num_id = None
    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()
        if not stripped:
            current_list_kind = None
            current_num_id = None
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_md_table(lines, idx)
            add_table(doc, rows)
            current_list_kind = None
            current_num_id = None
            continue

        if raw.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, raw[4:].strip(), base_size=12, base_color=DARK_BLUE, bold=True)
            current_list_kind = None
            current_num_id = None
            idx += 1
            continue
        if raw.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, raw[3:].strip(), base_size=13, base_color=BLUE, bold=True)
            current_list_kind = None
            current_num_id = None
            idx += 1
            continue
        if raw.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, raw[2:].strip(), base_size=16, base_color=BLUE, bold=True)
            current_list_kind = None
            current_num_id = None
            idx += 1
            continue

        match = LIST_RE.match(raw)
        if match:
            indent, marker, content = match.groups()
            kind = "bullet" if marker in ("-", "*") else "decimal"
            level = min(len(indent.replace("\t", "   ")) // 3, 3)
            if kind != current_list_kind or current_num_id is None:
                current_num_id = create_num_id(doc, kind)
                current_list_kind = kind
            p = doc.add_paragraph()
            apply_numbering(p, current_num_id, level)
            add_inline(p, content)
            idx += 1
            continue

        current_list_kind = None
        current_num_id = None
        if stripped.startswith("`") and stripped.endswith("`") and stripped.count("`") == 2:
            p = doc.add_paragraph(style="Formula Block")
            run = p.add_run(stripped[1:-1])
            set_run_font(run, name="Consolas", size=10, color=DARK_BLUE)
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        idx += 1


def build(source: Path, output: Path, kind: str):
    lines = source.read_text(encoding="utf-8").splitlines()
    title, meta, body_idx = split_front_matter(lines)
    doc = Document()
    configure_styles(doc)
    configure_section(doc, "Совет конституции голосования")
    doc.core_properties.title = title
    doc.core_properties.subject = "Проект институционального устройства тематического взвешенного голосования"
    doc.core_properties.keywords = "совет, экспертный вес, госплан, голосование, устав"

    if kind == "charter":
        add_memo_masthead(doc, title, meta)
    else:
        add_editorial_cover(doc, title, meta)
    add_markdown_body(doc, lines, body_idx)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--kind", choices=("charter", "explanation"), required=True)
    args = parser.parse_args()
    build(args.source, args.output, args.kind)


if __name__ == "__main__":
    main()
