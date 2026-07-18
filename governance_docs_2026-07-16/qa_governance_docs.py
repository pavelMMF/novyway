from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


CONTENT_DXA = 9360


def table_geometry(table):
    tbl = table._tbl
    grid = [int(c.get(qn("w:w"))) for c in tbl.tblGrid.findall(qn("w:gridCol"))]
    tbl_w = tbl.tblPr.find(qn("w:tblW"))
    tbl_ind = tbl.tblPr.find(qn("w:tblInd"))
    rows_ok = True
    for row in table.rows:
        widths = []
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            widths.append(int(tc_w.get(qn("w:w"))) if tc_w is not None else -1)
        rows_ok = rows_ok and widths == grid
    header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
    return {
        "grid": grid,
        "grid_sum": sum(grid),
        "tbl_width": int(tbl_w.get(qn("w:w"))) if tbl_w is not None else None,
        "indent": int(tbl_ind.get(qn("w:w"))) if tbl_ind is not None else None,
        "rows_match_grid": rows_ok,
        "header_flag": header,
    }


def inspect(path: Path, kind: str):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    errors = []
    warnings = []

    required = [
        "индивидуальный экспертный вес не имеет",
        "образовательная надбавка",
        "Госплан",
        "равный референдум" if kind == "explanation" else "Равный референдум",
    ]
    for phrase in required:
        if phrase.lower() not in text.lower():
            errors.append(f"missing required phrase: {phrase}")

    if re.search(r"\*\*|\[[^\]]+\]\(https?://|^\s*\|", text, re.M):
        errors.append("markdown residue found")
    if "TODO" in text or "PLACEHOLDER" in text:
        errors.append("placeholder found")

    headings = []
    for p in doc.paragraphs:
        m = re.fullmatch(r"Heading ([1-3])", p.style.name or "")
        if m:
            headings.append((int(m.group(1)), p.text))
    for idx in range(1, len(headings)):
        if headings[idx][0] > headings[idx - 1][0] + 1:
            errors.append(f"heading level skip: {headings[idx - 1]} -> {headings[idx]}")

    sec = doc.sections[0]
    geometry = {
        "page_width": round(sec.page_width.inches, 3),
        "page_height": round(sec.page_height.inches, 3),
        "margins": [round(x.inches, 3) for x in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin)],
        "header": round(sec.header_distance.inches, 3),
        "footer": round(sec.footer_distance.inches, 3),
    }
    if geometry["page_width"] != 8.5 or geometry["page_height"] != 11.0:
        errors.append(f"page geometry mismatch: {geometry}")
    if geometry["margins"] != [1.0, 1.0, 1.0, 1.0]:
        errors.append(f"margin mismatch: {geometry['margins']}")

    table_reports = [table_geometry(t) for t in doc.tables]
    for idx, report in enumerate(table_reports, start=1):
        if report["grid_sum"] != CONTENT_DXA or report["tbl_width"] != CONTENT_DXA:
            errors.append(f"table {idx} width mismatch: {report}")
        if report["indent"] != 120:
            errors.append(f"table {idx} indent mismatch: {report['indent']}")
        if not report["rows_match_grid"] or not report["header_flag"]:
            errors.append(f"table {idx} structural mismatch: {report}")

    num_count = 0
    for p in all_paragraphs:
        p_pr = p._p.pPr
        if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
            num_count += 1
    if num_count == 0:
        errors.append("no real numbered/bulleted list paragraphs")

    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        hyperlink_count = len(re.findall(r"TargetMode=\"External\"", rels))
        tracked = len(re.findall(r"<w:(?:ins|del)\b", document_xml))
        comments = any(name.startswith("word/comments") for name in names)
        if tracked or comments:
            errors.append(f"review residue: tracked={tracked}, comments={comments}")

    if kind == "explanation" and hyperlink_count < 10:
        errors.append(f"expected at least 10 source hyperlinks, found {hyperlink_count}")
    if kind == "charter" and len(doc.tables) < 1:
        errors.append("charter must contain composition table")
    if kind == "explanation" and len(doc.tables) < 1:
        errors.append("explanation must contain architecture comparison table")

    if kind == "charter":
        if "Общее число членов — 29" not in text:
            errors.append("College total 29 missing")
        if "индивидуальный экспертный вес не имеет числового верхнего предела" not in text:
            errors.append("uncapped expert invariant missing")
        if "универсальная совокупная доля экспертов" not in text.lower():
            errors.append("topic-specific quota invariant missing")
        if "девять представителей госплана" not in text.lower():
            errors.append("Gosplan 9-seat rule missing")

    return {
        "file": str(path),
        "kind": kind,
        "paragraphs": len(doc.paragraphs),
        "characters": len(text),
        "headings": len(headings),
        "tables": len(doc.tables),
        "numbered_or_bulleted_paragraphs": num_count,
        "external_hyperlinks": hyperlink_count,
        "page_geometry": geometry,
        "table_geometry": table_reports,
        "warnings": warnings,
        "errors": errors,
    }


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: qa_governance_docs.py <charter.docx> <explanation.docx>")
    reports = [inspect(Path(sys.argv[1]), "charter"), inspect(Path(sys.argv[2]), "explanation")]
    print(json.dumps(reports, ensure_ascii=False, indent=2))
    if any(r["errors"] for r in reports):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
